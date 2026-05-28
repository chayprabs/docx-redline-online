import base64
import io
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches
from lxml import etree

PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9l9XwAAAAASUVORK5CYII="
)


def create_conversion_fixture(target: Path) -> Path:
    image_path = target.with_suffix(".png")
    image_path.write_bytes(PNG_BYTES)

    document = Document()
    document.add_heading("DocxRedline Fixture", level=1)

    paragraph = document.add_paragraph()
    paragraph.add_run("This paragraph keeps ")
    paragraph.add_run("bold").bold = True
    paragraph.add_run(" and ")
    paragraph.add_run("italic").italic = True
    paragraph.add_run(" formatting.")

    document.add_paragraph("First bullet", style="List Bullet")
    document.add_paragraph("Second bullet", style="List Bullet")
    document.add_picture(str(image_path), width=Inches(1.0))

    document.save(target)
    return target


def create_comments_fixture(target: Path) -> Path:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Threaded ")
    target_run = paragraph.add_run("comment")
    document.add_comment(
        runs=[target_run],
        text="Top-level note",
        author="Ada",
        initials="AL",
    )
    document.save(target)

    with zipfile.ZipFile(target, "r") as archive:
        files = {name: archive.read(name) for name in archive.namelist()}

    comments_root = etree.fromstring(files["word/comments.xml"])
    comment_element = comments_root.xpath("./w:comment", namespaces={
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    })[0]
    first_paragraph = comment_element.xpath("./w:p", namespaces={
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    })[0]
    first_paragraph.set("{http://schemas.microsoft.com/office/word/2010/wordml}paraId", "00AAA111")

    reply = etree.Element(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment",
        nsmap=comments_root.nsmap,
    )
    reply.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", "1")
    reply.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author", "Ben")
    reply.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}initials", "BN")
    reply.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date", "2026-05-29T00:10:00Z")

    reply_paragraph = etree.SubElement(
        reply,
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p",
    )
    reply_paragraph.set("{http://schemas.microsoft.com/office/word/2010/wordml}paraId", "00BBB222")
    reply_run = etree.SubElement(
        reply_paragraph,
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r",
    )
    reply_text = etree.SubElement(
        reply_run,
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t",
    )
    reply_text.text = "Reply note"
    comments_root.append(reply)
    files["word/comments.xml"] = etree.tostring(
        comments_root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    comments_extended = etree.fromstring(
        b"""<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w15:commentsEx xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">
  <w15:commentEx w15:paraId="00AAA111" w15:done="1"/>
  <w15:commentEx w15:paraId="00BBB222" w15:paraIdParent="00AAA111"/>
</w15:commentsEx>
"""
    )
    files["word/commentsExtended.xml"] = etree.tostring(
        comments_extended,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    content_types = etree.fromstring(files["[Content_Types].xml"])
    override = etree.Element(
        "{http://schemas.openxmlformats.org/package/2006/content-types}Override"
    )
    override.set("PartName", "/word/commentsExtended.xml")
    override.set(
        "ContentType",
        "application/vnd.ms-word.commentsExtended+xml",
    )
    content_types.append(override)
    files["[Content_Types].xml"] = etree.tostring(
        content_types,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    relationships = etree.fromstring(files["word/_rels/document.xml.rels"])
    relationship = etree.Element(
        "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"
    )
    relationship.set("Id", "rIdCommentsEx1")
    relationship.set(
        "Type",
        "http://schemas.microsoft.com/office/2011/relationships/commentsExtended",
    )
    relationship.set("Target", "commentsExtended.xml")
    relationships.append(relationship)
    files["word/_rels/document.xml.rels"] = etree.tostring(
        relationships,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in files.items():
            archive.writestr(name, data)

    target.write_bytes(output.getvalue())
    return target


def create_tracked_changes_fixture(target: Path) -> Path:
    document = Document()
    document.add_paragraph("Tracked changes fixture")
    document.add_paragraph("Formatted paragraph")
    document.save(target)

    with zipfile.ZipFile(target, "r") as archive:
        files = {name: archive.read(name) for name in archive.namelist()}

    root = etree.fromstring(files["word/document.xml"])
    paragraphs = root.xpath("//w:body/w:p", namespaces={
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    })

    first_paragraph = paragraphs[0]
    for child in list(first_paragraph):
        first_paragraph.remove(child)

    run_start = etree.SubElement(first_paragraph, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
    text_start = etree.SubElement(run_start, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
    text_start.text = "Keep "

    deletion = etree.SubElement(first_paragraph, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}del")
    deletion.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", "1")
    deletion.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author", "Ada")
    deletion.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date", "2026-05-29T00:20:00Z")
    deletion_run = etree.SubElement(deletion, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
    deletion_text = etree.SubElement(deletion_run, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}delText")
    deletion_text.text = "old"

    insertion = etree.SubElement(first_paragraph, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ins")
    insertion.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", "2")
    insertion.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author", "Ben")
    insertion.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date", "2026-05-29T00:21:00Z")
    insertion_run = etree.SubElement(insertion, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
    insertion_text = etree.SubElement(insertion_run, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
    insertion_text.text = "new"

    move_from = etree.SubElement(first_paragraph, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}moveFrom")
    move_from.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", "3")
    move_from.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author", "Cara")
    move_from.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date", "2026-05-29T00:22:00Z")
    move_from_run = etree.SubElement(move_from, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
    move_from_text = etree.SubElement(move_from_run, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
    move_from_text.text = "away"

    move_to = etree.SubElement(first_paragraph, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}moveTo")
    move_to.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", "4")
    move_to.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author", "Cara")
    move_to.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date", "2026-05-29T00:22:00Z")
    move_to_run = etree.SubElement(move_to, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
    move_to_text = etree.SubElement(move_to_run, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
    move_to_text.text = "here"

    run_end = etree.SubElement(first_paragraph, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
    text_end = etree.SubElement(run_end, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
    text_end.text = " tail"

    second_paragraph = paragraphs[1]
    ppr = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    justification = etree.SubElement(ppr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc")
    justification.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "center")
    ppr_change = etree.SubElement(ppr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPrChange")
    ppr_change.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", "5")
    ppr_change.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author", "Dana")
    ppr_change.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date", "2026-05-29T00:23:00Z")
    old_ppr = etree.SubElement(ppr_change, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    old_justification = etree.SubElement(old_ppr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc")
    old_justification.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "left")
    second_paragraph.insert(0, ppr)

    files["word/document.xml"] = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in files.items():
            archive.writestr(name, data)

    target.write_bytes(output.getvalue())
    return target


def create_content_controls_fixture(target: Path) -> Path:
    document = Document()
    document.add_paragraph("Placeholder host")
    document.save(target)

    with zipfile.ZipFile(target, "r") as archive:
        files = {name: archive.read(name) for name in archive.namelist()}

    root = etree.fromstring(files["word/document.xml"])
    paragraph = root.xpath("//w:body/w:p", namespaces={
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    })[0]
    for child in list(paragraph):
        paragraph.remove(child)

    sdt = etree.SubElement(paragraph, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt")
    sdt_pr = etree.SubElement(sdt, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtPr")
    alias = etree.SubElement(sdt_pr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}alias")
    alias.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "Client Name")
    tag = etree.SubElement(sdt_pr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tag")
    tag.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "client_name")
    identifier = etree.SubElement(sdt_pr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
    identifier.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "101")
    etree.SubElement(sdt_pr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}text")

    sdt_content = etree.SubElement(sdt, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtContent")
    run_one = etree.SubElement(sdt_content, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
    run_one_properties = etree.SubElement(run_one, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
    etree.SubElement(run_one_properties, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}b")
    text_one = etree.SubElement(run_one, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
    text_one.text = "Acme"
    run_two = etree.SubElement(sdt_content, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
    text_two = etree.SubElement(run_two, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
    text_two.text = " Corp"

    files["word/document.xml"] = etree.tostring(
        root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in files.items():
            archive.writestr(name, data)

    target.write_bytes(output.getvalue())
    return target


def create_compare_fixtures(original_target: Path, revised_target: Path) -> tuple[Path, Path]:
    original = Document()
    original.add_paragraph("Alpha beta gamma")
    original.add_paragraph("Shared paragraph")
    original.save(original_target)

    revised = Document()
    revised.add_paragraph("Alpha delta gamma")
    revised.add_paragraph("Shared paragraph")
    revised.add_paragraph("New closing paragraph")
    revised.save(revised_target)

    return original_target, revised_target
