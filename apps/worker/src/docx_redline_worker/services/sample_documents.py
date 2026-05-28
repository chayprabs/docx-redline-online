import io
import os
import tempfile
import zipfile

from docx import Document
from docx.shared import Inches
from lxml import etree

PNG_BYTES = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x04"
    b"\x00\x00\x00\xb5\x1c\x0c\x02\x00\x00\x00\x0bIDATx\xdac\xfc\xff\x1f\x00\x02"
    b"\xeb\x01\xf5\x8fe\xf5|\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _document_bytes(document: Document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_contract_redline_sample(variant: str) -> bytes:
    document = Document()
    document.add_heading("Services Agreement", level=1)
    if variant == "original":
        document.add_paragraph("Payment is due within thirty days of invoice.")
        document.add_paragraph("Termination requires written notice from either party.")
    else:
        document.add_paragraph("Payment is due within fifteen days of invoice.")
        document.add_paragraph("Termination requires ten business days of written notice.")
        document.add_paragraph("Confidentiality obligations survive termination.")
    return _document_bytes(document)


def build_generic_images_sample() -> bytes:
    document = Document()
    document.add_heading("Illustrated Brief", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("This fixture keeps ")
    paragraph.add_run("bold").bold = True
    paragraph.add_run(" and ")
    paragraph.add_run("italic").italic = True
    paragraph.add_run(" formatting for conversion.")
    document.add_paragraph("First bullet", style="List Bullet")
    document.add_paragraph("Second bullet", style="List Bullet")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as handle:
        handle.write(PNG_BYTES)
        temp_name = handle.name
    try:
        document.add_picture(temp_name, width=Inches(1.0))
    finally:
        os.remove(temp_name)

    return _document_bytes(document)


def build_manuscript_comments_sample() -> bytes:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Threaded ")
    target_run = paragraph.add_run("comment")
    document.add_comment(
        runs=[target_run],
        text="Top-level note",
        author="Ada",
        initials="AL",
    )

    with zipfile.ZipFile(io.BytesIO(_document_bytes(document)), "r") as archive:
        files = {name: archive.read(name) for name in archive.namelist()}

    comments_root = etree.fromstring(files["word/comments.xml"])
    comment_element = comments_root.xpath(
        "./w:comment",
        namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
    )[0]
    first_paragraph = comment_element.xpath(
        "./w:p",
        namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
    )[0]
    first_paragraph.set(
        "{http://schemas.microsoft.com/office/word/2010/wordml}paraId",
        "00AAA111",
    )

    reply = etree.Element(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment",
        nsmap=comments_root.nsmap,
    )
    reply.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", "1")
    reply.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author", "Ben")
    reply.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}initials",
        "BN",
    )
    reply.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date",
        "2026-05-29T00:10:00Z",
    )
    reply_paragraph = etree.SubElement(
        reply,
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p",
    )
    reply_paragraph.set(
        "{http://schemas.microsoft.com/office/word/2010/wordml}paraId",
        "00BBB222",
    )
    reply_run = etree.SubElement(
        reply_paragraph,
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r",
    )
    reply_text = etree.SubElement(
        reply_run,
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t",
    )
    reply_text.text = "Reply note"
    comments_root.append(reply)
    files["word/comments.xml"] = etree.tostring(
        comments_root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    comments_extended = etree.fromstring(
        b"""<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w15:commentsEx xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">
  <w15:commentEx w15:paraId="00AAA111" w15:done="1"/>
  <w15:commentEx w15:paraId="00BBB222" w15:paraIdParent="00AAA111"/>
</w15:commentsEx>
"""
    )
    files["word/commentsExtended.xml"] = etree.tostring(
        comments_extended,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    content_types = etree.fromstring(files["[Content_Types].xml"])
    override = etree.Element(
        "{http://schemas.openxmlformats.org/package/2006/content-types}Override"
    )
    override.set("PartName", "/word/commentsExtended.xml")
    override.set("ContentType", "application/vnd.ms-word.commentsExtended+xml")
    content_types.append(override)
    files["[Content_Types].xml"] = etree.tostring(
        content_types,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    relationships = etree.fromstring(files["word/_rels/document.xml.rels"])
    relationship = etree.Element(
        "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"
    )
    relationship.set("Id", "rIdCommentsEx1")
    relationship.set(
        "Type",
        "http://schemas.microsoft.com/office/2011/relationships/commentsExtended",
    )
    relationship.set("Target", "commentsExtended.xml")
    relationships.append(relationship)
    files["word/_rels/document.xml.rels"] = etree.tostring(
        relationships,
        xml_declaration=True,
        encoding="UTF-8",
        standalone="yes",
    )

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in files.items():
            archive.writestr(name, data)

    return output.getvalue()


def get_sample_bytes(sample_id: str, variant: str | None = None) -> tuple[bytes, str]:
    if sample_id == "contract-redline":
        resolved_variant = variant or "original"
        if resolved_variant not in {"original", "revised"}:
            raise ValueError("contract-redline requires variant=original or variant=revised.")
        return (
            build_contract_redline_sample(resolved_variant),
            f"contract-redline-{resolved_variant}.docx",
        )

    if sample_id == "manuscript-comments":
        return build_manuscript_comments_sample(), "manuscript-comments.docx"

    if sample_id == "generic-images":
        return build_generic_images_sample(), "generic-images.docx"

    raise ValueError(f"Unknown sample '{sample_id}'.")
